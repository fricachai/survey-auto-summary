#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
survey_auto_summary.py

學術問卷資料前處理與自動分析摘要工具

主要功能
1. 問卷資料讀取、欄位檢核與反向計分
2. 缺失值與直線作答篩選
3. 題項與構面描述統計
4. 信度、KMO、Bartlett、EFA
5. 人口統計摘要
6. Pearson 相關分析
7. t 檢定與單因子 ANOVA
8. 多元迴歸與階層迴歸
9. 中介效果與調節效果分析
10. APA 風格表格與註腳
11. PLS-SEM 前置整理
12. 中文變項標籤對照
13. 自動圖表輸出與論文段落模板
"""

from __future__ import annotations

import argparse
import math
from pathlib import Path
from typing import Dict, Iterable, List, Sequence, Tuple

import numpy as np
import pandas as pd
from scipy import stats
import statsmodels.api as sm
from statsmodels.stats.outliers_influence import variance_inflation_factor

try:
    import matplotlib.pyplot as plt
except Exception:
    plt = None

try:
    from factor_analyzer import FactorAnalyzer
    from factor_analyzer.factor_analyzer import calculate_bartlett_sphericity, calculate_kmo
except Exception:
    FactorAnalyzer = None
    calculate_bartlett_sphericity = None
    calculate_kmo = None

try:
    from docx import Document
except Exception:
    Document = None


LIKERT_TYPE = "likert"
DEMOGRAPHIC_TYPE = "demographic"
CONTINUOUS_TYPE = "continuous"

ANALYSIS_REGRESSION = "regression"
ANALYSIS_HIERARCHICAL = "hierarchical_regression"
ANALYSIS_MEDIATION = "mediation"
ANALYSIS_MODERATION = "moderation"

CSV_ENCODINGS = ["utf-8-sig", "utf-8", "cp950", "big5"]
DEFAULT_ENCODING = "utf-8-sig"
BOOTSTRAP_SAMPLES = 1000

ITEM_MEAN_LOW_CUTOFF = 2.0
ITEM_MEAN_HIGH_CUTOFF = 4.0
ITEM_SD_LOW_CUTOFF = 0.50
ALPHA_LOW_CUTOFF = 0.70
LOW_ITEM_TOTAL_CUTOFF = 0.30
FACTOR_LOADING_LOW_CUTOFF = 0.50
FACTOR_CROSS_LOADING_CUTOFF = 0.40


def read_table(path: Path) -> pd.DataFrame:
    """讀取 CSV 或 Excel。"""
    if not path.exists():
        raise FileNotFoundError(f"找不到檔案：{path}")

    suffix = path.suffix.lower()
    if suffix in {".xlsx", ".xls"}:
        return pd.read_excel(path)
    if suffix == ".csv":
        last_error: Exception | None = None
        for encoding in CSV_ENCODINGS:
            try:
                return pd.read_csv(path, encoding=encoding)
            except Exception as exc:
                last_error = exc
        raise ValueError(f"CSV 讀取失敗：{path}；最後錯誤：{last_error}")
    raise ValueError(f"不支援的檔案格式：{suffix}")


def sanitize_sheet_name(name: str) -> str:
    invalid_chars = ["\\", "/", "*", "[", "]", ":", "?"]
    result = str(name).strip() or "Sheet1"
    for ch in invalid_chars:
        result = result.replace(ch, "_")
    return result[:31]


def write_excel(sheets: Dict[str, pd.DataFrame], output_path: Path) -> None:
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        for name, df in sheets.items():
            df.to_excel(writer, sheet_name=sanitize_sheet_name(name), index=False)


def ensure_required_columns(df: pd.DataFrame, columns: Sequence[str], table_name: str) -> None:
    missing = [col for col in columns if col not in df.columns]
    if missing:
        raise ValueError(f"{table_name} 缺少必要欄位：{', '.join(missing)}")


def normalize_config(config_df: pd.DataFrame) -> pd.DataFrame:
    ensure_required_columns(config_df, ["variable", "type"], "設定檔")
    config = config_df.copy()
    for col, default in {"construct": "", "subconstruct": "", "reverse": 0, "label": ""}.items():
        if col not in config.columns:
            config[col] = default
    config["variable"] = config["variable"].astype(str).str.strip()
    config["type"] = config["type"].astype(str).str.lower().str.strip()
    config["construct"] = config["construct"].fillna("").astype(str).str.strip()
    config["subconstruct"] = config["subconstruct"].fillna("").astype(str).str.strip()
    config["label"] = config["label"].fillna("").astype(str).str.strip()

    duplicated = config["variable"].duplicated(keep=False)
    if duplicated.any():
        dup_values = config.loc[duplicated, "variable"].tolist()
        raise ValueError(f"設定檔出現重複欄位名稱：{dup_values}")

    likert_rows = config["type"] == LIKERT_TYPE
    if likert_rows.any() and config.loc[likert_rows, "construct"].eq("").any():
        raise ValueError("設定檔中存在 type=likert 但 construct 為空白的列。")

    return config


def normalize_models(models_df: pd.DataFrame) -> pd.DataFrame:
    ensure_required_columns(models_df, ["analysis_type", "model_name", "outcome"], "模型設定檔")
    models = models_df.copy()
    for col in [
        "predictor",
        "mediator",
        "moderator",
        "covariates",
        "step_predictors",
        "notes",
    ]:
        if col not in models.columns:
            models[col] = ""
    models["analysis_type"] = models["analysis_type"].astype(str).str.lower().str.strip()
    for col in ["model_name", "outcome", "predictor", "mediator", "moderator", "covariates", "step_predictors", "notes"]:
        models[col] = models[col].fillna("").astype(str).str.strip()
    return models


def normalize_labels(labels_df: pd.DataFrame) -> Dict[str, str]:
    ensure_required_columns(labels_df, ["variable", "label"], "標籤對照檔")
    labels = labels_df.copy()
    labels["variable"] = labels["variable"].astype(str).str.strip()
    labels["label"] = labels["label"].astype(str).str.strip()
    labels = labels.loc[labels["variable"] != ""]
    return dict(zip(labels["variable"], labels["label"]))


def build_label_map(config: pd.DataFrame, extra_labels: Dict[str, str] | None) -> Dict[str, str]:
    label_map: Dict[str, str] = {}
    for _, row in config.iterrows():
        label = row.get("label", "")
        if label:
            label_map[str(row["variable"])] = str(label)
    if extra_labels:
        label_map.update(extra_labels)
    return label_map


def label_for(value: str, label_map: Dict[str, str]) -> str:
    return label_map.get(str(value), str(value))


def relabel_value_column(df: pd.DataFrame, column: str, label_map: Dict[str, str]) -> pd.DataFrame:
    if column not in df.columns:
        return df
    out = df.copy()
    out[column] = out[column].map(lambda x: label_for(x, label_map))
    return out


def parse_var_list(value: object) -> List[str]:
    if value is None or pd.isna(value):
        return []
    text = str(value).strip()
    if not text:
        return []
    return [part.strip() for part in text.split("|") if part.strip()]


def parse_step_predictors(value: object) -> List[List[str]]:
    if value is None or pd.isna(value):
        return []
    text = str(value).strip()
    if not text:
        return []
    return [parse_var_list(step) for step in text.split("||") if step.strip()]


def coerce_numeric_columns(df: pd.DataFrame, columns: Iterable[str]) -> pd.DataFrame:
    out = df.copy()
    for col in columns:
        if col in out.columns:
            out[col] = pd.to_numeric(out[col], errors="coerce")
    return out


def p_label(p_value: float) -> str:
    if pd.isna(p_value):
        return "NA"
    if p_value < 0.001:
        return "p < .001"
    return f"p = {p_value:.3f}".replace("0.", ".")


def significance_star(p_value: float) -> str:
    if pd.isna(p_value):
        return ""
    if p_value < 0.001:
        return "***"
    if p_value < 0.01:
        return "**"
    if p_value < 0.05:
        return "*"
    return ""


def cronbach_alpha(df: pd.DataFrame) -> float:
    clean = df.dropna(axis=0, how="any")
    if clean.shape[0] < 2 or clean.shape[1] < 2:
        return np.nan
    item_var = clean.var(axis=0, ddof=1)
    total_var = clean.sum(axis=1).var(ddof=1)
    if total_var == 0:
        return np.nan
    n_items = clean.shape[1]
    return float((n_items / (n_items - 1)) * (1 - item_var.sum() / total_var))


def corrected_item_total_statistics(df: pd.DataFrame) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    clean = df.dropna(axis=0, how="any")
    for col in df.columns:
        if clean.shape[0] < 3 or df.shape[1] < 2:
            citc = np.nan
        else:
            total_minus_item = clean.drop(columns=[col]).sum(axis=1)
            if clean[col].std(ddof=1) == 0 or total_minus_item.std(ddof=1) == 0:
                citc = np.nan
            else:
                citc = clean[col].corr(total_minus_item)
        rows.append(
            {
                "題項": col,
                "校正後題項總分相關": citc,
                "刪題後_alpha": cronbach_alpha(df.drop(columns=[col])),
            }
        )
    return pd.DataFrame(rows)


def descriptive_statistics(df: pd.DataFrame, label_name: str) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame(columns=[label_name, "有效樣本數", "平均數", "標準差", "最小值", "最大值", "中位數"])
    return pd.DataFrame(
        {
            label_name: df.columns,
            "有效樣本數": df.notna().sum().values,
            "平均數": df.mean(numeric_only=True).values,
            "標準差": df.std(numeric_only=True, ddof=1).values,
            "最小值": df.min(numeric_only=True).values,
            "最大值": df.max(numeric_only=True).values,
            "中位數": df.median(numeric_only=True).values,
        }
    )


def ceiling_floor_table(df: pd.DataFrame, scale_min: int, scale_max: int, threshold: float = 0.5) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    for col in df.columns:
        s = df[col].dropna()
        floor_ratio = (s == scale_min).mean() if not s.empty else np.nan
        ceiling_ratio = (s == scale_max).mean() if not s.empty else np.nan
        rows.append(
            {
                "題項": col,
                "地板比例": floor_ratio,
                "天花板比例": ceiling_ratio,
                "地板警示": bool(pd.notna(floor_ratio) and floor_ratio >= threshold),
                "天花板警示": bool(pd.notna(ceiling_ratio) and ceiling_ratio >= threshold),
            }
        )
    return pd.DataFrame(rows)


def choose_factor_count(df: pd.DataFrame) -> Tuple[int, pd.DataFrame]:
    corr = df.dropna().corr()
    eigvals, _ = np.linalg.eig(corr.values)
    eigvals = np.real(eigvals)
    eigvals = np.sort(eigvals)[::-1]
    count = max(1, int((eigvals > 1).sum()))
    eig_df = pd.DataFrame({"因素": [f"F{i+1}" for i in range(len(eigvals))], "特徵值": eigvals})
    return count, eig_df


def run_kmo_bartlett(df: pd.DataFrame) -> Tuple[float, float, float]:
    if calculate_kmo is None or calculate_bartlett_sphericity is None:
        return np.nan, np.nan, np.nan
    clean = df.dropna()
    if clean.shape[0] < 5 or clean.shape[1] < 2:
        return np.nan, np.nan, np.nan
    _, kmo_model = calculate_kmo(clean)
    chi_square_value, p_value = calculate_bartlett_sphericity(clean)
    return float(kmo_model), float(chi_square_value), float(p_value)


def run_efa(df: pd.DataFrame, rotation: str = "varimax") -> Tuple[pd.DataFrame, pd.DataFrame]:
    if FactorAnalyzer is None:
        return pd.DataFrame(), pd.DataFrame()
    clean = df.dropna()
    if clean.shape[0] < 5 or clean.shape[1] < 2:
        return pd.DataFrame(), pd.DataFrame()
    n_factors, eigen_df = choose_factor_count(clean)
    try:
        analyzer = FactorAnalyzer(n_factors=n_factors, rotation=rotation)
        analyzer.fit(clean)
    except Exception:
        return pd.DataFrame(), eigen_df
    loading_df = pd.DataFrame(
        analyzer.loadings_,
        index=clean.columns,
        columns=[f"F{i+1}" for i in range(n_factors)],
    ).reset_index().rename(columns={"index": "題項"})
    return loading_df, eigen_df


def mean_sd_label(mean_value: float, sd_value: float) -> str:
    if pd.isna(mean_value) or pd.isna(sd_value):
        return "NA"
    return f"{mean_value:.2f} ({sd_value:.2f})"


def reverse_score_items(df: pd.DataFrame, reverse_vars: Sequence[str], scale_min: int, scale_max: int) -> pd.DataFrame:
    out = df.copy()
    for col in reverse_vars:
        if col in out.columns:
            out[col] = np.where(out[col].notna(), scale_min + scale_max - out[col], np.nan)
    return out


def build_range_issue_table(df: pd.DataFrame, item_vars: Sequence[str], scale_min: int, scale_max: int) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    for col in item_vars:
        count = int(((df[col].notna()) & ((df[col] < scale_min) | (df[col] > scale_max))).sum())
        if count > 0:
            rows.append({"題項": col, "異常筆數": count, "說明": f"共有 {count} 筆超出量表範圍 {scale_min} 至 {scale_max}。"})
    return pd.DataFrame(rows)


def screen_cases(df: pd.DataFrame, item_vars: Sequence[str], missing_threshold: float, straightline_check: bool) -> Tuple[pd.DataFrame, pd.DataFrame]:
    out = df.copy()
    if item_vars:
        out["__缺失比例過高__"] = out[list(item_vars)].isna().mean(axis=1) > missing_threshold
        out["__直線作答__"] = out[list(item_vars)].nunique(axis=1, dropna=True) <= 1 if straightline_check else False
    else:
        out["__缺失比例過高__"] = False
        out["__直線作答__"] = False
    clean = out.loc[~out["__缺失比例過高__"] & ~out["__直線作答__"]].copy().reset_index(drop=True)
    summary = pd.DataFrame(
        [
            {"指標": "原始樣本數", "數值": len(out)},
            {"指標": "有效樣本數", "數值": len(clean)},
            {"指標": "因缺失比例過高刪除", "數值": int(out["__缺失比例過高__"].sum())},
            {"指標": "因直線作答刪除", "數值": int(out["__直線作答__"].sum())},
        ]
    )
    return clean, summary


def build_construct_scores(clean_df: pd.DataFrame, likert_config: pd.DataFrame) -> pd.DataFrame:
    if likert_config.empty:
        return pd.DataFrame()
    scores: Dict[str, pd.Series] = {}
    for construct, sub_df in likert_config.groupby("construct"):
        scores[str(construct)] = clean_df[sub_df["variable"].tolist()].mean(axis=1)
    score_df = pd.DataFrame(scores)
    score_df.insert(0, "受試者編號", range(1, len(score_df) + 1))
    return score_df


def build_construct_reliability_outputs(clean_df: pd.DataFrame, likert_config: pd.DataFrame) -> Tuple[pd.DataFrame, Dict[str, pd.DataFrame], pd.DataFrame, Dict[str, pd.DataFrame], Dict[str, pd.DataFrame]]:
    alpha_rows: List[Dict[str, object]] = []
    kmo_rows: List[Dict[str, object]] = []
    item_total_tables: Dict[str, pd.DataFrame] = {}
    loading_tables: Dict[str, pd.DataFrame] = {}
    eigen_tables: Dict[str, pd.DataFrame] = {}

    for construct, sub_df in likert_config.groupby("construct"):
        items = sub_df["variable"].tolist()
        cdf = clean_df[items].copy()
        alpha_rows.append(
            {
                "構面": construct,
                "題數": len(items),
                "有效樣本數": int(cdf.dropna().shape[0]),
                "Cronbach_alpha": cronbach_alpha(cdf),
            }
        )
        item_total_tables[str(construct)] = corrected_item_total_statistics(cdf)
        kmo_value, chi_square, p_value = run_kmo_bartlett(cdf)
        kmo_rows.append({"構面": construct, "KMO": kmo_value, "Bartlett卡方值": chi_square, "Bartlett_p值": p_value})
        loading_df, eigen_df = run_efa(cdf)
        if not loading_df.empty:
            loading_tables[str(construct)] = loading_df
        if not eigen_df.empty:
            eigen_tables[str(construct)] = eigen_df

    return pd.DataFrame(alpha_rows), item_total_tables, pd.DataFrame(kmo_rows), loading_tables, eigen_tables


def build_demographic_summary(df: pd.DataFrame, demographic_vars: Sequence[str], label_map: Dict[str, str]) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    for variable in demographic_vars:
        series = df[variable]
        counts = series.value_counts(dropna=False)
        valid_n = int(series.notna().sum())
        cumulative = 0.0
        for category, count in counts.items():
            is_missing = pd.isna(category)
            total_pct = count / len(series) * 100 if len(series) else np.nan
            valid_pct = np.nan if is_missing or valid_n == 0 else count / valid_n * 100
            if pd.notna(valid_pct):
                cumulative += valid_pct
            rows.append(
                {
                    "人口統計變項": label_for(variable, label_map),
                    "原始欄位": variable,
                    "類別": "遺漏值" if is_missing else str(category),
                    "次數": int(count),
                    "百分比": total_pct,
                    "有效百分比": valid_pct,
                    "累積有效百分比": cumulative if pd.notna(valid_pct) else np.nan,
                }
            )
    return pd.DataFrame(rows)


def build_correlation_results(construct_score_df: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    numeric_df = construct_score_df.drop(columns=["受試者編號"], errors="ignore")
    if numeric_df.shape[1] < 2:
        return pd.DataFrame(columns=["構面"]), pd.DataFrame(columns=["變項一", "變項二", "有效樣本數", "Pearson_r", "p值", "顯著性"])

    cols = numeric_df.columns.tolist()
    matrix = pd.DataFrame("", index=cols, columns=cols)
    detail_rows: List[Dict[str, object]] = []
    for i, col_x in enumerate(cols):
        matrix.loc[col_x, col_x] = "1.00"
        for col_y in cols[i + 1 :]:
            sub = numeric_df[[col_x, col_y]].dropna()
            if len(sub) < 3:
                label = "NA"
                r_value = np.nan
                p_value = np.nan
            else:
                r_value, p_value = stats.pearsonr(sub[col_x], sub[col_y])
                label = f"{r_value:.3f}{significance_star(p_value)}"
                detail_rows.append(
                    {
                        "變項一": col_x,
                        "變項二": col_y,
                        "有效樣本數": len(sub),
                        "Pearson_r": r_value,
                        "p值": p_value,
                        "顯著性": significance_star(p_value),
                    }
                )
            matrix.loc[col_x, col_y] = label
            matrix.loc[col_y, col_x] = label
    return matrix.reset_index().rename(columns={"index": "構面"}), pd.DataFrame(detail_rows)


def cohen_d(group1: pd.Series, group2: pd.Series) -> float:
    n1, n2 = len(group1), len(group2)
    if n1 < 2 or n2 < 2:
        return np.nan
    s1, s2 = group1.std(ddof=1), group2.std(ddof=1)
    pooled = np.sqrt(((n1 - 1) * s1**2 + (n2 - 1) * s2**2) / (n1 + n2 - 2))
    if pooled == 0:
        return np.nan
    return float((group1.mean() - group2.mean()) / pooled)


def eta_squared(groups: Sequence[pd.Series]) -> float:
    all_values = pd.concat(groups, ignore_index=True)
    if all_values.empty:
        return np.nan
    grand_mean = all_values.mean()
    ss_between = sum(len(g) * (g.mean() - grand_mean) ** 2 for g in groups)
    ss_total = ((all_values - grand_mean) ** 2).sum()
    if ss_total == 0:
        return np.nan
    return float(ss_between / ss_total)


def build_group_difference_results(clean_df: pd.DataFrame, demographic_vars: Sequence[str], construct_score_df: pd.DataFrame, label_map: Dict[str, str]) -> Tuple[pd.DataFrame, pd.DataFrame]:
    construct_df = construct_score_df.drop(columns=["受試者編號"], errors="ignore")
    if construct_df.empty:
        return pd.DataFrame(), pd.DataFrame()

    analysis_df = pd.concat([clean_df[list(demographic_vars)].copy(), construct_df], axis=1)
    t_rows: List[Dict[str, object]] = []
    a_rows: List[Dict[str, object]] = []
    for group_var in demographic_vars:
        for construct in construct_df.columns:
            sub = analysis_df[[group_var, construct]].dropna()
            if sub.empty:
                continue
            levels = list(pd.unique(sub[group_var]))
            groups = [sub.loc[sub[group_var] == level, construct].astype(float) for level in levels]
            groups = [g for g in groups if len(g) > 0]
            if len(groups) < 2:
                continue
            if len(levels) == 2:
                g1, g2 = groups
                t_stat, p_value = stats.ttest_ind(g1, g2, equal_var=False, nan_policy="omit")
                df_num = ((g1.var(ddof=1) / len(g1) + g2.var(ddof=1) / len(g2)) ** 2)
                df_den = 0.0
                if len(g1) > 1:
                    df_den += ((g1.var(ddof=1) / len(g1)) ** 2) / (len(g1) - 1)
                if len(g2) > 1:
                    df_den += ((g2.var(ddof=1) / len(g2)) ** 2) / (len(g2) - 1)
                df_value = df_num / df_den if df_den else np.nan
                t_rows.append(
                    {
                        "分組變項": label_for(group_var, label_map),
                        "構面": construct,
                        "組別一": str(levels[0]),
                        "組別一樣本數": len(g1),
                        "組別一平均數": g1.mean(),
                        "組別一標準差": g1.std(ddof=1),
                        "組別二": str(levels[1]),
                        "組別二樣本數": len(g2),
                        "組別二平均數": g2.mean(),
                        "組別二標準差": g2.std(ddof=1),
                        "t值": t_stat,
                        "自由度": df_value,
                        "p值": p_value,
                        "效果量_d": cohen_d(g1, g2),
                    }
                )
            elif len(levels) >= 3:
                if any(len(g) < 2 for g in groups):
                    a_rows.append(
                        {
                            "分組變項": label_for(group_var, label_map),
                            "構面": construct,
                            "組數": len(groups),
                            "總樣本數": sum(len(g) for g in groups),
                            "F值": np.nan,
                            "組間自由度": len(groups) - 1,
                            "組內自由度": sum(len(g) for g in groups) - len(groups),
                            "p值": np.nan,
                            "效果量_eta平方": np.nan,
                            "備註": "至少一組有效樣本數不足 2，未執行 ANOVA",
                        }
                    )
                    continue
                f_stat, p_value = stats.f_oneway(*groups)
                a_rows.append(
                    {
                        "分組變項": label_for(group_var, label_map),
                        "構面": construct,
                        "組數": len(groups),
                        "總樣本數": sum(len(g) for g in groups),
                        "F值": f_stat,
                        "組間自由度": len(groups) - 1,
                        "組內自由度": sum(len(g) for g in groups) - len(groups),
                        "p值": p_value,
                        "效果量_eta平方": eta_squared(groups),
                        "備註": "",
                    }
                )
    return pd.DataFrame(t_rows), pd.DataFrame(a_rows)


def zscore_df(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for col in out.columns:
        s = out[col]
        sd = s.std(ddof=1)
        if pd.isna(sd) or sd == 0:
            out[col] = np.nan
        else:
            out[col] = (s - s.mean()) / sd
    return out


def standardized_beta(y: pd.Series, x_df: pd.DataFrame) -> pd.Series:
    z_y = zscore_df(pd.DataFrame({"y": y}))["y"]
    z_x = zscore_df(x_df)
    if z_x.dropna().empty or z_y.dropna().empty:
        return pd.Series(dtype=float)
    model = sm.OLS(z_y, sm.add_constant(z_x, has_constant="add"), missing="drop").fit()
    return model.params.drop("const", errors="ignore")


def fit_ols(df: pd.DataFrame, outcome: str, predictors: Sequence[str]) -> Tuple[sm.regression.linear_model.RegressionResultsWrapper | None, pd.DataFrame]:
    needed = [outcome, *predictors]
    sub = df[needed].dropna().copy()
    if len(sub) < max(8, len(predictors) + 3):
        return None, sub
    y = sub[outcome]
    x = sm.add_constant(sub[list(predictors)], has_constant="add")
    try:
        model = sm.OLS(y, x).fit()
    except Exception:
        return None, sub
    return model, sub


def build_coefficients_table(model, sub_df: pd.DataFrame, outcome: str, label_map: Dict[str, str], model_name: str, step_name: str = "") -> pd.DataFrame:
    std_beta = standardized_beta(sub_df[outcome], sub_df[[col for col in sub_df.columns if col != outcome]])
    rows: List[Dict[str, object]] = []
    for term in model.params.index:
        beta_value = np.nan
        if term != "const" and term in std_beta.index:
            beta_value = std_beta[term]
        rows.append(
            {
                "模型名稱": model_name,
                "步驟": step_name,
                "依變項": label_for(outcome, label_map),
                "項目": "常數" if term == "const" else label_for(term, label_map),
                "B": model.params[term],
                "SE": model.bse[term],
                "Beta": beta_value,
                "t值": model.tvalues[term],
                "p值": model.pvalues[term],
            }
        )
    return pd.DataFrame(rows)


def run_regression_models(analysis_df: pd.DataFrame, models_df: pd.DataFrame, label_map: Dict[str, str]) -> Tuple[pd.DataFrame, pd.DataFrame]:
    rows: List[Dict[str, object]] = []
    coef_tables: List[pd.DataFrame] = []
    model_df = models_df.loc[models_df["analysis_type"] == ANALYSIS_REGRESSION].copy()
    for _, row in model_df.iterrows():
        predictors = parse_var_list(row["predictor"]) + parse_var_list(row["covariates"])
        outcome = row["outcome"]
        model_name = row["model_name"]
        missing = [var for var in [outcome, *predictors] if var not in analysis_df.columns]
        if missing or not predictors:
            rows.append(
                {
                    "模型名稱": model_name,
                    "依變項": label_for(outcome, label_map),
                    "自變項": " | ".join(label_for(var, label_map) for var in predictors),
                    "樣本數": np.nan,
                    "R平方": np.nan,
                    "調整後R平方": np.nan,
                    "F值": np.nan,
                    "p值": np.nan,
                    "備註": f"缺少欄位：{', '.join(missing)}" if missing else "未提供自變項",
                }
            )
            continue
        model, sub = fit_ols(analysis_df, outcome, predictors)
        if model is None:
            rows.append(
                {
                    "模型名稱": model_name,
                    "依變項": label_for(outcome, label_map),
                    "自變項": " | ".join(label_for(var, label_map) for var in predictors),
                    "樣本數": len(sub),
                    "R平方": np.nan,
                    "調整後R平方": np.nan,
                    "F值": np.nan,
                    "p值": np.nan,
                    "備註": "模型估計失敗或有效樣本不足",
                }
            )
            continue
        rows.append(
            {
                "模型名稱": model_name,
                "依變項": label_for(outcome, label_map),
                "自變項": " | ".join(label_for(var, label_map) for var in predictors),
                "樣本數": int(model.nobs),
                "R平方": model.rsquared,
                "調整後R平方": model.rsquared_adj,
                "F值": model.fvalue,
                "p值": model.f_pvalue,
                "備註": row["notes"],
            }
        )
        coef_tables.append(build_coefficients_table(model, sub, outcome, label_map, model_name))
    coef_df = pd.concat(coef_tables, ignore_index=True) if coef_tables else pd.DataFrame()
    return pd.DataFrame(rows), coef_df


def f_change_test(r2_prev: float, r2_curr: float, n: int, p_prev: int, p_curr: int) -> Tuple[float, float]:
    if n <= p_curr + 1 or p_curr <= p_prev:
        return np.nan, np.nan
    numerator = (r2_curr - r2_prev) / (p_curr - p_prev)
    denominator = (1 - r2_curr) / (n - p_curr - 1)
    if denominator == 0:
        return np.nan, np.nan
    f_change = numerator / denominator
    p_value = stats.f.sf(f_change, p_curr - p_prev, n - p_curr - 1)
    return float(f_change), float(p_value)


def run_hierarchical_models(analysis_df: pd.DataFrame, models_df: pd.DataFrame, label_map: Dict[str, str]) -> Tuple[pd.DataFrame, pd.DataFrame]:
    rows: List[Dict[str, object]] = []
    coef_tables: List[pd.DataFrame] = []
    model_df = models_df.loc[models_df["analysis_type"] == ANALYSIS_HIERARCHICAL].copy()
    for _, row in model_df.iterrows():
        steps = parse_step_predictors(row["step_predictors"])
        outcome = row["outcome"]
        model_name = row["model_name"]
        if not steps:
            rows.append(
                {
                    "模型名稱": model_name,
                    "步驟": "",
                    "依變項": label_for(outcome, label_map),
                    "自變項": "",
                    "樣本數": np.nan,
                    "R平方": np.nan,
                    "調整後R平方": np.nan,
                    "R平方變化量": np.nan,
                    "F變化量": np.nan,
                    "F變化量_p值": np.nan,
                    "備註": "未設定 step_predictors",
                }
            )
            continue
        current_predictors: List[str] = []
        prev_r2 = 0.0
        prev_p = 0
        for step_index, step_vars in enumerate(steps, start=1):
            current_predictors.extend(step_vars)
            missing = [var for var in [outcome, *current_predictors] if var not in analysis_df.columns]
            step_name = f"Step {step_index}"
            if missing:
                rows.append(
                    {
                        "模型名稱": model_name,
                        "步驟": step_name,
                        "依變項": label_for(outcome, label_map),
                        "自變項": " | ".join(label_for(var, label_map) for var in current_predictors),
                        "樣本數": np.nan,
                        "R平方": np.nan,
                        "調整後R平方": np.nan,
                        "R平方變化量": np.nan,
                        "F變化量": np.nan,
                        "F變化量_p值": np.nan,
                        "備註": f"缺少欄位：{', '.join(missing)}",
                    }
                )
                continue
            model, sub = fit_ols(analysis_df, outcome, current_predictors)
            if model is None:
                rows.append(
                    {
                        "模型名稱": model_name,
                        "步驟": step_name,
                        "依變項": label_for(outcome, label_map),
                        "自變項": " | ".join(label_for(var, label_map) for var in current_predictors),
                        "樣本數": len(sub),
                        "R平方": np.nan,
                        "調整後R平方": np.nan,
                        "R平方變化量": np.nan,
                        "F變化量": np.nan,
                        "F變化量_p值": np.nan,
                        "備註": "模型估計失敗或有效樣本不足",
                    }
                )
                continue
            curr_p = len(current_predictors)
            delta_r2 = model.rsquared - prev_r2 if step_index > 1 else model.rsquared
            f_change, p_change = f_change_test(prev_r2, model.rsquared, int(model.nobs), prev_p, curr_p) if step_index > 1 else (np.nan, np.nan)
            rows.append(
                {
                    "模型名稱": model_name,
                    "步驟": step_name,
                    "依變項": label_for(outcome, label_map),
                    "自變項": " | ".join(label_for(var, label_map) for var in current_predictors),
                    "樣本數": int(model.nobs),
                    "R平方": model.rsquared,
                    "調整後R平方": model.rsquared_adj,
                    "R平方變化量": delta_r2,
                    "F變化量": f_change,
                    "F變化量_p值": p_change,
                    "備註": row["notes"],
                }
            )
            coef_tables.append(build_coefficients_table(model, sub, outcome, label_map, model_name, step_name))
            prev_r2 = model.rsquared
            prev_p = curr_p
    coef_df = pd.concat(coef_tables, ignore_index=True) if coef_tables else pd.DataFrame()
    return pd.DataFrame(rows), coef_df


def bootstrap_indirect_effect(df: pd.DataFrame, x_var: str, m_var: str, y_var: str, covariates: Sequence[str], seed: int = 42) -> Tuple[float, float]:
    rng = np.random.default_rng(seed)
    effects: List[float] = []
    base = df[[x_var, m_var, y_var, *covariates]].dropna().copy()
    if len(base) < 5:
        return np.nan, np.nan
    for _ in range(BOOTSTRAP_SAMPLES):
        sample = base.iloc[rng.integers(0, len(base), len(base))].copy()
        a_model, _ = fit_ols(sample, m_var, [x_var, *covariates])
        b_model, _ = fit_ols(sample, y_var, [x_var, m_var, *covariates])
        if a_model is None or b_model is None:
            continue
        effects.append(a_model.params.get(x_var, np.nan) * b_model.params.get(m_var, np.nan))
    if not effects:
        return np.nan, np.nan
    lower = float(np.nanpercentile(effects, 2.5))
    upper = float(np.nanpercentile(effects, 97.5))
    return lower, upper


def run_mediation_models(analysis_df: pd.DataFrame, models_df: pd.DataFrame, label_map: Dict[str, str]) -> Tuple[pd.DataFrame, pd.DataFrame]:
    summary_rows: List[Dict[str, object]] = []
    detail_tables: List[pd.DataFrame] = []
    model_df = models_df.loc[models_df["analysis_type"] == ANALYSIS_MEDIATION].copy()
    for _, row in model_df.iterrows():
        x_var = row["predictor"]
        m_var = row["mediator"]
        y_var = row["outcome"]
        covariates = parse_var_list(row["covariates"])
        model_name = row["model_name"]
        missing = [var for var in [x_var, m_var, y_var, *covariates] if var not in analysis_df.columns]
        if missing or not x_var or not m_var:
            summary_rows.append(
                {
                    "模型名稱": model_name,
                    "自變項X": label_for(x_var, label_map),
                    "中介變項M": label_for(m_var, label_map),
                    "依變項Y": label_for(y_var, label_map),
                    "樣本數": np.nan,
                    "a路徑": np.nan,
                    "b路徑": np.nan,
                    "c總效果": np.nan,
                    "c'直接效果": np.nan,
                    "間接效果ab": np.nan,
                    "Sobel_z": np.nan,
                    "Sobel_p值": np.nan,
                    "Bootstrap_CI下限": np.nan,
                    "Bootstrap_CI上限": np.nan,
                    "備註": f"缺少欄位：{', '.join(missing)}" if missing else "中介模型規格不完整",
                }
            )
            continue
        a_model, a_sub = fit_ols(analysis_df, m_var, [x_var, *covariates])
        c_model, c_sub = fit_ols(analysis_df, y_var, [x_var, *covariates])
        b_model, b_sub = fit_ols(analysis_df, y_var, [x_var, m_var, *covariates])
        if a_model is None or c_model is None or b_model is None:
            summary_rows.append(
                {
                    "模型名稱": model_name,
                    "自變項X": label_for(x_var, label_map),
                    "中介變項M": label_for(m_var, label_map),
                    "依變項Y": label_for(y_var, label_map),
                    "樣本數": np.nan,
                    "a路徑": np.nan,
                    "b路徑": np.nan,
                    "c總效果": np.nan,
                    "c'直接效果": np.nan,
                    "間接效果ab": np.nan,
                    "Sobel_z": np.nan,
                    "Sobel_p值": np.nan,
                    "Bootstrap_CI下限": np.nan,
                    "Bootstrap_CI上限": np.nan,
                    "備註": "模型估計失敗或有效樣本不足",
                }
            )
            continue
        a = a_model.params.get(x_var, np.nan)
        b = b_model.params.get(m_var, np.nan)
        c_total = c_model.params.get(x_var, np.nan)
        c_prime = b_model.params.get(x_var, np.nan)
        a_se = a_model.bse.get(x_var, np.nan)
        b_se = b_model.bse.get(m_var, np.nan)
        indirect = a * b
        sobel_se = math.sqrt((b**2) * (a_se**2) + (a**2) * (b_se**2)) if pd.notna(a) and pd.notna(b) and pd.notna(a_se) and pd.notna(b_se) else np.nan
        sobel_z = indirect / sobel_se if pd.notna(sobel_se) and sobel_se != 0 else np.nan
        sobel_p = 2 * stats.norm.sf(abs(sobel_z)) if pd.notna(sobel_z) else np.nan
        ci_low, ci_high = bootstrap_indirect_effect(analysis_df, x_var, m_var, y_var, covariates)
        summary_rows.append(
            {
                "模型名稱": model_name,
                "自變項X": label_for(x_var, label_map),
                "中介變項M": label_for(m_var, label_map),
                "依變項Y": label_for(y_var, label_map),
                "樣本數": min(len(a_sub), len(b_sub)),
                "a路徑": a,
                "b路徑": b,
                "c總效果": c_total,
                "c'直接效果": c_prime,
                "間接效果ab": indirect,
                "Sobel_z": sobel_z,
                "Sobel_p值": sobel_p,
                "Bootstrap_CI下限": ci_low,
                "Bootstrap_CI上限": ci_high,
                "備註": row["notes"],
            }
        )
        detail_tables.append(build_coefficients_table(a_model, a_sub, m_var, label_map, model_name, "a路徑"))
        detail_tables.append(build_coefficients_table(c_model, c_sub, y_var, label_map, model_name, "c總效果"))
        detail_tables.append(build_coefficients_table(b_model, b_sub, y_var, label_map, model_name, "b與c'"))
    detail_df = pd.concat(detail_tables, ignore_index=True) if detail_tables else pd.DataFrame()
    return pd.DataFrame(summary_rows), detail_df


def run_moderation_models(analysis_df: pd.DataFrame, models_df: pd.DataFrame, label_map: Dict[str, str]) -> Tuple[pd.DataFrame, pd.DataFrame]:
    summary_rows: List[Dict[str, object]] = []
    coef_tables: List[pd.DataFrame] = []
    model_df = models_df.loc[models_df["analysis_type"] == ANALYSIS_MODERATION].copy()
    for _, row in model_df.iterrows():
        x_var = row["predictor"]
        w_var = row["moderator"]
        y_var = row["outcome"]
        covariates = parse_var_list(row["covariates"])
        model_name = row["model_name"]
        missing = [var for var in [x_var, w_var, y_var, *covariates] if var not in analysis_df.columns]
        if missing or not x_var or not w_var:
            summary_rows.append(
                {
                    "模型名稱": model_name,
                    "自變項X": label_for(x_var, label_map),
                    "調節變項W": label_for(w_var, label_map),
                    "依變項Y": label_for(y_var, label_map),
                    "樣本數": np.nan,
                    "R平方": np.nan,
                    "調整後R平方": np.nan,
                    "交互作用B": np.nan,
                    "交互作用Beta": np.nan,
                    "交互作用t值": np.nan,
                    "交互作用p值": np.nan,
                    "備註": f"缺少欄位：{', '.join(missing)}" if missing else "調節模型規格不完整",
                }
            )
            continue
        sub = analysis_df[[x_var, w_var, y_var, *covariates]].dropna().copy()
        if len(sub) < max(8, len(covariates) + 5):
            summary_rows.append(
                {
                    "模型名稱": model_name,
                    "自變項X": label_for(x_var, label_map),
                    "調節變項W": label_for(w_var, label_map),
                    "依變項Y": label_for(y_var, label_map),
                    "樣本數": len(sub),
                    "R平方": np.nan,
                    "調整後R平方": np.nan,
                    "交互作用B": np.nan,
                    "交互作用Beta": np.nan,
                    "交互作用t值": np.nan,
                    "交互作用p值": np.nan,
                    "備註": "有效樣本不足",
                }
            )
            continue
        x_center = f"{x_var}_centered"
        w_center = f"{w_var}_centered"
        interaction = f"{x_var}_X_{w_var}"
        sub[x_center] = sub[x_var] - sub[x_var].mean()
        sub[w_center] = sub[w_var] - sub[w_var].mean()
        sub[interaction] = sub[x_center] * sub[w_center]
        predictors = [x_center, w_center, interaction, *covariates]
        model, fit_sub = fit_ols(sub, y_var, predictors)
        if model is None:
            summary_rows.append(
                {
                    "模型名稱": model_name,
                    "自變項X": label_for(x_var, label_map),
                    "調節變項W": label_for(w_var, label_map),
                    "依變項Y": label_for(y_var, label_map),
                    "樣本數": len(sub),
                    "R平方": np.nan,
                    "調整後R平方": np.nan,
                    "交互作用B": np.nan,
                    "交互作用Beta": np.nan,
                    "交互作用t值": np.nan,
                    "交互作用p值": np.nan,
                    "備註": "模型估計失敗",
                }
            )
            continue
        coef_df = build_coefficients_table(model, fit_sub, y_var, label_map, model_name)
        coef_df["項目"] = coef_df["項目"].replace(
            {
                label_for(x_center, label_map): f"{label_for(x_var, label_map)}（中心化）",
                label_for(w_center, label_map): f"{label_for(w_var, label_map)}（中心化）",
                label_for(interaction, label_map): f"{label_for(x_var, label_map)} × {label_for(w_var, label_map)}",
                x_center: f"{label_for(x_var, label_map)}（中心化）",
                w_center: f"{label_for(w_var, label_map)}（中心化）",
                interaction: f"{label_for(x_var, label_map)} × {label_for(w_var, label_map)}",
            }
        )
        coef_tables.append(coef_df)
        interaction_row = coef_df.loc[coef_df["項目"] == f"{label_for(x_var, label_map)} × {label_for(w_var, label_map)}"]
        if interaction_row.empty:
            interaction_row = coef_df.tail(1)
        ir = interaction_row.iloc[0]
        summary_rows.append(
            {
                "模型名稱": model_name,
                "自變項X": label_for(x_var, label_map),
                "調節變項W": label_for(w_var, label_map),
                "依變項Y": label_for(y_var, label_map),
                "樣本數": int(model.nobs),
                "R平方": model.rsquared,
                "調整後R平方": model.rsquared_adj,
                "交互作用B": ir["B"],
                "交互作用Beta": ir["Beta"],
                "交互作用t值": ir["t值"],
                "交互作用p值": ir["p值"],
                "備註": row["notes"],
            }
        )
    coef_df = pd.concat(coef_tables, ignore_index=True) if coef_tables else pd.DataFrame()
    return pd.DataFrame(summary_rows), coef_df


def compute_vif(df: pd.DataFrame, label_map: Dict[str, str]) -> pd.DataFrame:
    numeric_df = df.dropna().copy()
    if numeric_df.shape[1] < 2:
        return pd.DataFrame(columns=["變項", "VIF"])
    x = sm.add_constant(numeric_df, has_constant="add")
    rows = []
    for index, col in enumerate(x.columns):
        if col == "const":
            continue
        try:
            vif_value = variance_inflation_factor(x.values, index)
        except Exception:
            vif_value = np.nan
        rows.append({"變項": label_for(col, label_map), "VIF": vif_value})
    return pd.DataFrame(rows)


def build_pls_sem_outputs(clean_df: pd.DataFrame, likert_config: pd.DataFrame, construct_score_df: pd.DataFrame, label_map: Dict[str, str]) -> Dict[str, pd.DataFrame]:
    if likert_config.empty:
        return {
            "PLS_SEM前置說明": pd.DataFrame([{"說明": "未設定量表題項，故未產生 PLS-SEM 前置整理資料。"}])
        }
    item_vars = likert_config["variable"].tolist()
    indicator_df = clean_df[item_vars].copy()
    indicator_df.insert(0, "受試者編號", range(1, len(indicator_df) + 1))
    z_indicator = zscore_df(indicator_df.drop(columns=["受試者編號"], errors="ignore"))
    z_indicator.insert(0, "受試者編號", indicator_df["受試者編號"])
    mapping_df = likert_config[["variable", "construct", "subconstruct", "reverse"]].copy()
    mapping_df["題項標籤"] = mapping_df["variable"].map(lambda x: label_for(x, label_map))
    mapping_df = mapping_df.rename(
        columns={"variable": "題項", "construct": "構面", "subconstruct": "子構面", "reverse": "反向計分"}
    )
    missing_sd_df = pd.DataFrame(
        {
            "題項": item_vars,
            "題項標籤": [label_for(item, label_map) for item in item_vars],
            "缺失比例": clean_df[item_vars].isna().mean().values,
            "平均數": clean_df[item_vars].mean().values,
            "標準差": clean_df[item_vars].std(ddof=1).values,
        }
    )
    construct_numeric = construct_score_df.drop(columns=["受試者編號"], errors="ignore")
    construct_vif_df = compute_vif(construct_numeric, label_map)
    construct_z = zscore_df(construct_numeric)
    construct_z.insert(0, "受試者編號", construct_score_df["受試者編號"])
    notes_df = pd.DataFrame(
        [
            {"說明": "PLS_題項資料：可作為 SmartPLS 或其他軟體之前置題項資料。"},
            {"說明": "PLS_Z標準化題項：提供標準化後題項值，便於外部模型軟體測試。"},
            {"說明": "PLS_測量模型對照：提供題項、構面、子構面與反向計分對照。"},
            {"說明": "PLS_構面Z分數：提供構面平均分數標準化後結果。"},
            {"說明": "PLS_VIF檢核：提供構面層級之共線性參考。"},
        ]
    )
    return {
        "PLS_題項資料": indicator_df,
        "PLS_Z標準化題項": z_indicator,
        "PLS_測量模型對照": mapping_df,
        "PLS_缺失與變異檢核": missing_sd_df,
        "PLS_構面Z分數": construct_z,
        "PLS_VIF檢核": construct_vif_df,
        "PLS_SEM前置說明": notes_df,
    }


def build_warning_table(item_desc_df: pd.DataFrame, construct_alpha_df: pd.DataFrame, item_total_tables: Dict[str, pd.DataFrame], loading_tables: Dict[str, pd.DataFrame], range_issue_df: pd.DataFrame) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    for _, row in item_desc_df.iterrows():
        item_name = row["題項"]
        mean_value = row["平均數"]
        sd_value = row["標準差"]
        if pd.notna(mean_value) and mean_value >= ITEM_MEAN_HIGH_CUTOFF:
            rows.append({"層級": "題項", "名稱": item_name, "警示類型": "平均數偏高", "說明": f"題項平均數為 {mean_value:.3f}，高於門檻 {ITEM_MEAN_HIGH_CUTOFF:.2f}。"})
        if pd.notna(mean_value) and mean_value <= ITEM_MEAN_LOW_CUTOFF:
            rows.append({"層級": "題項", "名稱": item_name, "警示類型": "平均數偏低", "說明": f"題項平均數為 {mean_value:.3f}，低於門檻 {ITEM_MEAN_LOW_CUTOFF:.2f}。"})
        if pd.notna(sd_value) and sd_value <= ITEM_SD_LOW_CUTOFF:
            rows.append({"層級": "題項", "名稱": item_name, "警示類型": "標準差偏低", "說明": f"題項標準差為 {sd_value:.3f}，低於門檻 {ITEM_SD_LOW_CUTOFF:.2f}。"})
        if bool(row.get("天花板警示", False)):
            rows.append({"層級": "題項", "名稱": item_name, "警示類型": "天花板效果", "說明": f"題項天花板比例為 {row['天花板比例']:.3f}。"})
        if bool(row.get("地板警示", False)):
            rows.append({"層級": "題項", "名稱": item_name, "警示類型": "地板效果", "說明": f"題項地板比例為 {row['地板比例']:.3f}。"})

    for _, row in construct_alpha_df.iterrows():
        alpha_value = row["Cronbach_alpha"]
        if pd.notna(alpha_value) and alpha_value < ALPHA_LOW_CUTOFF:
            rows.append({"層級": "構面", "名稱": row["構面"], "警示類型": "信度偏低", "說明": f"Cronbach's alpha 為 {alpha_value:.3f}，低於 {ALPHA_LOW_CUTOFF:.2f}。"})

    for construct_name, table in item_total_tables.items():
        current_alpha_series = construct_alpha_df.loc[construct_alpha_df["構面"] == construct_name, "Cronbach_alpha"]
        current_alpha = current_alpha_series.iloc[0] if not current_alpha_series.empty else np.nan
        for _, row in table.iterrows():
            citc = row["校正後題項總分相關"]
            alpha_deleted = row["刪題後_alpha"]
            if pd.notna(citc) and citc < LOW_ITEM_TOTAL_CUTOFF:
                rows.append({"層級": "題項", "名稱": row["題項"], "警示類型": "題項總分相關偏低", "說明": f"{construct_name} 題項 {row['題項']} 之 CITC = {citc:.3f}。"})
            if pd.notna(alpha_deleted) and pd.notna(current_alpha) and alpha_deleted > current_alpha + 0.02:
                rows.append({"層級": "題項", "名稱": row["題項"], "警示類型": "刪題後信度上升", "說明": f"{construct_name} 若刪除 {row['題項']}，alpha 可由 {current_alpha:.3f} 提升至 {alpha_deleted:.3f}。"})

    for construct_name, loading_df in loading_tables.items():
        factor_cols = [col for col in loading_df.columns if col.startswith("F")]
        for _, row in loading_df.iterrows():
            values = [abs(row[col]) for col in factor_cols if pd.notna(row[col])]
            if not values:
                continue
            largest = max(values)
            second = sorted(values, reverse=True)[1] if len(values) > 1 else 0
            if largest < FACTOR_LOADING_LOW_CUTOFF:
                rows.append({"層級": "題項", "名稱": row["題項"], "警示類型": "因素負荷量偏低", "說明": f"{construct_name} 題項 {row['題項']} 最大因素負荷量為 {largest:.3f}。"})
            if second >= FACTOR_CROSS_LOADING_CUTOFF:
                rows.append({"層級": "題項", "名稱": row["題項"], "警示類型": "交叉負荷偏高", "說明": f"{construct_name} 題項 {row['題項']} 次高負荷量為 {second:.3f}。"})

    if not range_issue_df.empty:
        for _, row in range_issue_df.iterrows():
            rows.append({"層級": "題項", "名稱": row["題項"], "警示類型": "超出量表範圍", "說明": row["說明"]})

    return pd.DataFrame(rows) if rows else pd.DataFrame(columns=["層級", "名稱", "警示類型", "說明"])


def build_package_notes() -> List[str]:
    notes: List[str] = []
    if FactorAnalyzer is None:
        notes.append("factor_analyzer 未安裝：已略過 KMO、Bartlett 與 EFA。")
    if Document is None:
        notes.append("python-docx 未安裝：已略過 Word 報告輸出。")
    if plt is None:
        notes.append("matplotlib 未安裝：已略過圖表輸出。")
    return notes


def build_project_check(root: Path) -> pd.DataFrame:
    expected = [
        "survey_auto_summary.py",
        "questionnaire_config_template.csv",
        "sample_data_template.csv",
        "analysis_model_template.csv",
        "sample_variable_labels.csv",
        "README_survey_auto_summary.md",
        "requirements.txt",
        "run.bat",
    ]
    return pd.DataFrame(
        [
            {"檔案名稱": name, "是否存在": (root / name).exists(), "大小_位元組": (root / name).stat().st_size if (root / name).exists() else np.nan}
            for name in expected
        ]
    )


def build_apa_demographic_table(demo_df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    if demo_df.empty:
        return pd.DataFrame(columns=["人口統計變項", "類別", "n", "%", "有效%", "累積有效%"]), "註：未設定人口統計變項或無有效資料。"
    out = demo_df.rename(columns={"次數": "n", "百分比": "%", "有效百分比": "有效%", "累積有效百分比": "累積有效%"})
    return out[["人口統計變項", "類別", "n", "%", "有效%", "累積有效%"]], "註：百分比以原始樣本數計算；有效百分比與累積有效百分比以非遺漏樣本計算。"


def build_apa_construct_desc_table(construct_desc_df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    if construct_desc_df.empty:
        return pd.DataFrame(columns=["構面", "n", "M", "SD", "Min", "Max"]), "註：未產生構面描述統計。"
    out = construct_desc_df.rename(columns={"有效樣本數": "n", "平均數": "M", "標準差": "SD", "最小值": "Min", "最大值": "Max"})
    return out[["構面", "n", "M", "SD", "Min", "Max"]], "註：構面分數以各題平均分數計算。"


def build_apa_ttest_table(t_df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    if t_df.empty:
        return pd.DataFrame(columns=["分組變項", "構面", "組別一 M(SD)", "組別二 M(SD)", "t(df)", "p", "d"]), "註：無可執行之雙組比較。"
    out = t_df.copy()
    out["組別一 M(SD)"] = out.apply(lambda r: mean_sd_label(r["組別一平均數"], r["組別一標準差"]), axis=1)
    out["組別二 M(SD)"] = out.apply(lambda r: mean_sd_label(r["組別二平均數"], r["組別二標準差"]), axis=1)
    out["t(df)"] = out.apply(lambda r: f"{r['t值']:.2f} ({r['自由度']:.2f})" if pd.notna(r["t值"]) and pd.notna(r["自由度"]) else "NA", axis=1)
    out["p"] = out["p值"].apply(p_label)
    out["d"] = out["效果量_d"].map(lambda v: "NA" if pd.isna(v) else f"{v:.2f}")
    return out[["分組變項", "構面", "組別一 M(SD)", "組別二 M(SD)", "t(df)", "p", "d"]], "註：採 Welch t 檢定；d 為 Cohen's d。"


def build_apa_anova_table(anova_df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    if anova_df.empty:
        return pd.DataFrame(columns=["分組變項", "構面", "F(df1, df2)", "p", "eta^2", "備註"]), "註：無可執行之單因子 ANOVA。"
    out = anova_df.copy()
    out["F(df1, df2)"] = out.apply(lambda r: f"{r['F值']:.2f} ({int(r['組間自由度'])}, {int(r['組內自由度'])})" if pd.notna(r["F值"]) else "NA", axis=1)
    out["p"] = out["p值"].apply(p_label)
    out["eta^2"] = out["效果量_eta平方"].map(lambda v: "NA" if pd.isna(v) else f"{v:.2f}")
    return out[["分組變項", "構面", "F(df1, df2)", "p", "eta^2", "備註"]], "註：eta^2 為效果量指標。"


def build_apa_correlation_table(corr_detail_df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    if corr_detail_df.empty:
        return pd.DataFrame(columns=["變項一", "變項二", "n", "r", "p"]), "註：無可執行之 Pearson 相關分析。"
    out = corr_detail_df.copy()
    out["n"] = out["有效樣本數"]
    out["r"] = out["Pearson_r"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["p"] = out["p值"].apply(p_label)
    return out[["變項一", "變項二", "n", "r", "p"]], "註：* p < .05，** p < .01，*** p < .001。"


def build_apa_regression_table(reg_df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    if reg_df.empty:
        return pd.DataFrame(columns=["模型名稱", "依變項", "自變項", "n", "R^2", "Adj. R^2", "F", "p", "備註"]), "註：未設定或未成功估計多元迴歸模型。"
    out = reg_df.copy()
    out["n"] = out["樣本數"]
    out["R^2"] = out["R平方"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["Adj. R^2"] = out["調整後R平方"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["F"] = out["F值"].map(lambda v: "NA" if pd.isna(v) else f"{v:.2f}")
    out["p"] = out["p值"].apply(p_label)
    return out[["模型名稱", "依變項", "自變項", "n", "R^2", "Adj. R^2", "F", "p", "備註"]], "註：本表為多元迴歸整體模型摘要。"


def build_apa_hierarchical_table(hier_df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    if hier_df.empty:
        return pd.DataFrame(columns=["模型名稱", "步驟", "依變項", "自變項", "R^2", "Delta R^2", "F change", "p", "備註"]), "註：未設定或未成功估計階層迴歸模型。"
    out = hier_df.copy()
    out["R^2"] = out["R平方"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["Delta R^2"] = out["R平方變化量"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["F change"] = out["F變化量"].map(lambda v: "NA" if pd.isna(v) else f"{v:.2f}")
    out["p"] = out["F變化量_p值"].apply(p_label)
    return out[["模型名稱", "步驟", "依變項", "自變項", "R^2", "Delta R^2", "F change", "p", "備註"]], "註：Delta R^2 表示相較前一步驟之解釋力變化。"


def build_apa_mediation_table(med_df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    if med_df.empty:
        return pd.DataFrame(columns=["模型名稱", "X", "M", "Y", "a", "b", "c", "c'", "ab", "Sobel z", "Sobel p", "95% CI", "備註"]), "註：未設定或未成功估計中介模型。"
    out = med_df.copy()
    out["a"] = out["a路徑"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["b"] = out["b路徑"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["c"] = out["c總效果"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["c'"] = out["c'直接效果"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["ab"] = out["間接效果ab"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["Sobel z"] = out["Sobel_z"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["Sobel p"] = out["Sobel_p值"].apply(p_label)
    out["95% CI"] = out.apply(
        lambda r: f"[{r['Bootstrap_CI下限']:.3f}, {r['Bootstrap_CI上限']:.3f}]" if pd.notna(r["Bootstrap_CI下限"]) and pd.notna(r["Bootstrap_CI上限"]) else "NA",
        axis=1,
    )
    return out.rename(columns={"自變項X": "X", "中介變項M": "M", "依變項Y": "Y"})[
        ["模型名稱", "X", "M", "Y", "a", "b", "c", "c'", "ab", "Sobel z", "Sobel p", "95% CI", "備註"]
    ], "註：95% CI 為 bootstrap percentile 信賴區間；若區間不含 0，可視為具中介效果參考。"


def build_apa_moderation_table(mod_df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    if mod_df.empty:
        return pd.DataFrame(columns=["模型名稱", "X", "W", "Y", "R^2", "Adj. R^2", "Interaction B", "Interaction Beta", "t", "p", "備註"]), "註：未設定或未成功估計調節模型。"
    out = mod_df.copy()
    out["R^2"] = out["R平方"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["Adj. R^2"] = out["調整後R平方"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["Interaction B"] = out["交互作用B"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["Interaction Beta"] = out["交互作用Beta"].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}")
    out["t"] = out["交互作用t值"].map(lambda v: "NA" if pd.isna(v) else f"{v:.2f}")
    out["p"] = out["交互作用p值"].apply(p_label)
    return out.rename(columns={"自變項X": "X", "調節變項W": "W", "依變項Y": "Y"})[
        ["模型名稱", "X", "W", "Y", "R^2", "Adj. R^2", "Interaction B", "Interaction Beta", "t", "p", "備註"]
    ], "註：交互作用項顯著時，可進一步進行 simple slope 檢驗。"


def build_apa_coeff_table(coef_df: pd.DataFrame) -> pd.DataFrame:
    if coef_df.empty:
        return pd.DataFrame(columns=["模型名稱", "步驟", "依變項", "項目", "B", "SE", "Beta", "t值", "p值"])
    out = coef_df.copy()
    for col in ["B", "SE", "Beta", "t值"]:
        out[col] = out[col].map(lambda v: "NA" if pd.isna(v) else f"{v:.3f}" if col != "t值" else f"{v:.2f}")
    out["p值"] = out["p值"].apply(p_label)
    return out[["模型名稱", "步驟", "依變項", "項目", "B", "SE", "Beta", "t值", "p值"]]


def add_word_table(doc: Document, title: str, df: pd.DataFrame, note: str | None = None) -> None:
    doc.add_heading(title, level=2)
    if df.empty:
        doc.add_paragraph("本節無可輸出結果。")
        if note:
            doc.add_paragraph(f"註：{note}")
        return
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if pd.isna(value):
                cells[i].text = ""
            elif isinstance(value, float):
                cells[i].text = f"{value:.3f}" if abs(value) < 1000 else f"{value:.2f}"
            else:
                cells[i].text = str(value)
    if note:
        doc.add_paragraph(f"註：{note}")


def build_summary_text(sample_info_df: pd.DataFrame, demo_df: pd.DataFrame, alpha_df: pd.DataFrame, kmo_df: pd.DataFrame, corr_df: pd.DataFrame, t_df: pd.DataFrame, anova_df: pd.DataFrame, reg_df: pd.DataFrame, hier_df: pd.DataFrame, med_df: pd.DataFrame, mod_df: pd.DataFrame, warning_df: pd.DataFrame) -> str:
    total_n = int(sample_info_df.loc[sample_info_df["指標"] == "原始樣本數", "數值"].iloc[0])
    valid_n = int(sample_info_df.loc[sample_info_df["指標"] == "有效樣本數", "數值"].iloc[0])
    removed_missing = int(sample_info_df.loc[sample_info_df["指標"] == "因缺失比例過高刪除", "數值"].iloc[0])
    removed_straight = int(sample_info_df.loc[sample_info_df["指標"] == "因直線作答刪除", "數值"].iloc[0])
    parts = [
        f"本研究原始資料共 {total_n} 份。經資料清理後，因缺失比例過高刪除 {removed_missing} 份，因直線作答刪除 {removed_straight} 份，最終有效樣本為 {valid_n} 份。",
    ]
    if not demo_df.empty:
        demo_sentences = []
        for variable, sub_df in demo_df.groupby("人口統計變項"):
            valid_sub = sub_df[sub_df["類別"] != "遺漏值"]
            if not valid_sub.empty:
                top = valid_sub.sort_values("次數", ascending=False).iloc[0]
                demo_sentences.append(f"在「{variable}」變項中，以「{top['類別']}」最多，占有效樣本 {top['有效百分比']:.1f}%。")
        if demo_sentences:
            parts.append("".join(demo_sentences))
    if not alpha_df.empty and alpha_df["Cronbach_alpha"].dropna().any():
        alpha_min = alpha_df["Cronbach_alpha"].min()
        alpha_max = alpha_df["Cronbach_alpha"].max()
        parts.append(f"各構面 Cronbach's alpha 介於 {alpha_min:.3f} 至 {alpha_max:.3f}。")
    if not kmo_df.empty and kmo_df["KMO"].dropna().any():
        parts.append(f"KMO 值介於 {kmo_df['KMO'].dropna().min():.3f} 至 {kmo_df['KMO'].dropna().max():.3f}。")
    if not corr_df.empty:
        sig_count = int((corr_df["p值"] < 0.05).fillna(False).sum())
        parts.append(f"Pearson 相關分析共完成 {len(corr_df)} 組配對，其中達顯著者 {sig_count} 組。")
    parts.append(f"雙組比較完成 {len(t_df)} 項，ANOVA 完成 {len(anova_df)} 項。")
    parts.append(f"多元迴歸模型 {len(reg_df)} 個，階層迴歸步驟 {len(hier_df)} 個，中介模型 {len(med_df)} 個，調節模型 {len(mod_df)} 個。")
    parts.append(f"資料品質警示共 {len(warning_df)} 項，研究者宜進一步檢視相關題項與構面。")
    return "\n\n".join(parts)


def build_thesis_paragraphs(summary_text: str, demo_df: pd.DataFrame, alpha_df: pd.DataFrame, kmo_df: pd.DataFrame, corr_detail_df: pd.DataFrame, t_df: pd.DataFrame, anova_df: pd.DataFrame, reg_df: pd.DataFrame, hier_df: pd.DataFrame, med_df: pd.DataFrame, mod_df: pd.DataFrame) -> str:
    lines = [
        "【樣本與資料清理段落】",
        summary_text,
        "",
        "【人口統計描述段落】",
        "本研究首先針對樣本之人口統計背景進行描述統計，以了解受試者組成結構。" if not demo_df.empty else "本研究未納入人口統計描述變項。",
        "",
        "【信度分析段落】",
        "本研究以 Cronbach's alpha 檢驗各構面之內部一致性。" if not alpha_df.empty else "本研究未能產出信度分析結果。",
        "",
        "【效度分析段落】",
        "本研究以 KMO 與 Bartlett 球形檢定作為因素分析適切性之參考，並視資料條件進一步檢視探索性因素分析結果。" if not kmo_df.empty else "本研究未能產出完整效度分析結果。",
        "",
        "【相關分析段落】",
        f"本研究進一步檢驗各構面間之 Pearson 積差相關，共完成 {len(corr_detail_df)} 組配對分析。" if not corr_detail_df.empty else "本研究未進行 Pearson 相關分析。",
        "",
        "【差異分析段落】",
        f"針對不同背景變項，本研究分別進行 t 檢定 {len(t_df)} 項與單因子變異數分析 {len(anova_df)} 項，以了解不同群體在構面上的差異情形。",
        "",
        "【多元迴歸段落】",
        f"本研究共估計 {len(reg_df)} 個多元迴歸模型，用以檢視自變項對依變項之解釋力。" if not reg_df.empty else "本研究未設定多元迴歸模型。",
        "",
        "【階層迴歸段落】",
        f"本研究以階層迴歸方式逐步投入解釋變項，共完成 {len(hier_df)} 個步驟之估計。" if not hier_df.empty else "本研究未設定階層迴歸模型。",
        "",
        "【中介效果段落】",
        f"本研究共檢驗 {len(med_df)} 個中介效果模型，並同時參考 Sobel 檢定與 bootstrap 信賴區間。" if not med_df.empty else "本研究未設定中介效果模型。",
        "",
        "【調節效果段落】",
        f"本研究共檢驗 {len(mod_df)} 個調節效果模型，以交互作用項檢驗調節效果是否成立。" if not mod_df.empty else "本研究未設定調節效果模型。",
    ]
    return "\n".join(lines)


def create_charts(output_dir: Path, demo_df: pd.DataFrame, construct_desc_df: pd.DataFrame, corr_matrix_df: pd.DataFrame) -> Tuple[pd.DataFrame, str]:
    if plt is None:
        return pd.DataFrame(columns=["圖表檔名", "圖表類型", "說明"]), "未安裝 matplotlib，已略過圖表輸出。"
    chart_dir = output_dir / "charts"
    chart_dir.mkdir(parents=True, exist_ok=True)
    plt.rcParams["font.sans-serif"] = ["Microsoft JhengHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    rows: List[Dict[str, object]] = []
    for variable, sub_df in demo_df.groupby("人口統計變項"):
        valid_sub = sub_df[sub_df["類別"] != "遺漏值"]
        if valid_sub.empty:
            continue
        fig, ax = plt.subplots(figsize=(7, 4))
        ax.bar(valid_sub["類別"], valid_sub["次數"], color="#4C78A8")
        ax.set_title(f"{variable} 分布圖")
        ax.set_xlabel("類別")
        ax.set_ylabel("次數")
        plt.xticks(rotation=30)
        fig.tight_layout()
        file_path = chart_dir / f"{variable}_bar.png"
        fig.savefig(file_path, dpi=200)
        plt.close(fig)
        rows.append({"圖表檔名": file_path.name, "圖表類型": "人口統計長條圖", "說明": f"{variable} 類別分布圖"})

    if not construct_desc_df.empty:
        fig, ax = plt.subplots(figsize=(7, 4))
        ax.bar(construct_desc_df["構面"], construct_desc_df["平均數"], color="#72B7B2")
        ax.set_title("構面平均數比較圖")
        ax.set_xlabel("構面")
        ax.set_ylabel("平均數")
        plt.xticks(rotation=30)
        fig.tight_layout()
        file_path = chart_dir / "construct_means_bar.png"
        fig.savefig(file_path, dpi=200)
        plt.close(fig)
        rows.append({"圖表檔名": file_path.name, "圖表類型": "構面平均數長條圖", "說明": "各構面平均數比較圖"})

    if not corr_matrix_df.empty and "構面" in corr_matrix_df.columns:
        corr_plot_df = corr_matrix_df.set_index("構面").replace("NA", np.nan)
        numeric_matrix = corr_plot_df.apply(lambda col: col.map(lambda x: float(str(x).replace("*", "")) if pd.notna(x) and str(x) not in ["", "NA"] else np.nan))
        fig, ax = plt.subplots(figsize=(6, 5))
        im = ax.imshow(numeric_matrix.values, cmap="Blues", vmin=-1, vmax=1)
        ax.set_xticks(range(len(numeric_matrix.columns)))
        ax.set_xticklabels(numeric_matrix.columns, rotation=45, ha="right")
        ax.set_yticks(range(len(numeric_matrix.index)))
        ax.set_yticklabels(numeric_matrix.index)
        ax.set_title("構面相關熱圖")
        fig.colorbar(im, ax=ax)
        fig.tight_layout()
        file_path = chart_dir / "construct_correlation_heatmap.png"
        fig.savefig(file_path, dpi=200)
        plt.close(fig)
        rows.append({"圖表檔名": file_path.name, "圖表類型": "相關熱圖", "說明": "構面相關矩陣熱圖"})

    return pd.DataFrame(rows), ""


def write_word_report(output_path: Path, summary_text: str, apa_tables: List[Tuple[str, pd.DataFrame, str]], chart_df: pd.DataFrame, package_notes: List[str]) -> str:
    if Document is None:
        return "未安裝 python-docx，已略過 Word 報告輸出。"
    doc = Document()
    doc.add_heading("問卷資料自動分析摘要報告", level=1)
    doc.add_paragraph(summary_text)
    if package_notes:
        doc.add_heading("執行環境備註", level=2)
        for note in package_notes:
            doc.add_paragraph(note, style="List Bullet")
    for title, table, note in apa_tables:
        add_word_table(doc, title, table, note)
    if not chart_df.empty:
        add_word_table(doc, "圖表輸出清單", chart_df, "註：圖表檔案另存於 output/charts 資料夾。")
    doc.save(output_path)
    return ""


def build_output_notes(package_notes: List[str], messages: Sequence[str]) -> pd.DataFrame:
    rows = [{"備註": note} for note in package_notes]
    rows.extend({"備註": message} for message in messages if message)
    return pd.DataFrame(rows) if rows else pd.DataFrame(columns=["備註"])


def relabel_construct_outputs(construct_score_df: pd.DataFrame, label_map: Dict[str, str]) -> pd.DataFrame:
    if construct_score_df.empty:
        return construct_score_df
    out = construct_score_df.copy()
    rename_map = {col: label_for(col, label_map) for col in out.columns if col != "受試者編號"}
    return out.rename(columns=rename_map)


def apply_term_labels_to_df(df: pd.DataFrame, columns: Sequence[str], label_map: Dict[str, str]) -> pd.DataFrame:
    out = df.copy()
    for col in columns:
        if col in out.columns:
            out[col] = out[col].map(lambda x: label_for(x, label_map))
    return out


def parse_arguments() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="學術問卷資料前處理與自動分析摘要工具")
    parser.add_argument("--data", required=True, help="問卷資料檔路徑（CSV / Excel）")
    parser.add_argument("--config", required=True, help="問卷設定檔路徑（CSV / Excel）")
    parser.add_argument("--models", help="模型設定檔路徑（CSV / Excel），用於迴歸 / 階層迴歸 / 中介 / 調節")
    parser.add_argument("--labels", help="中文變項標籤對照檔路徑（CSV / Excel）")
    parser.add_argument("--outdir", required=True, help="輸出資料夾")
    parser.add_argument("--scale-min", type=int, default=1, help="Likert 最小值，預設 1")
    parser.add_argument("--scale-max", type=int, default=5, help="Likert 最大值，預設 5")
    parser.add_argument("--missing-threshold", type=float, default=0.2, help="個案缺失比例刪除門檻，預設 0.2")
    parser.add_argument("--straightline-check", action="store_true", help="是否檢查直線作答")
    return parser.parse_args()


def main() -> None:
    args = parse_arguments()
    package_notes = build_package_notes()

    try:
        outdir = Path(args.outdir)
        outdir.mkdir(parents=True, exist_ok=True)

        data_df = read_table(Path(args.data))
        config_df = normalize_config(read_table(Path(args.config)))

        labels_map = {}
        if args.labels:
            labels_map = normalize_labels(read_table(Path(args.labels)))
        label_map = build_label_map(config_df, labels_map)

        ensure_required_columns(data_df, config_df["variable"].tolist(), "資料檔")

        likert_config = config_df.loc[config_df["type"] == LIKERT_TYPE].copy()
        demographic_config = config_df.loc[config_df["type"] == DEMOGRAPHIC_TYPE].copy()
        continuous_config = config_df.loc[config_df["type"] == CONTINUOUS_TYPE].copy()

        item_vars = likert_config["variable"].tolist()
        demographic_vars = demographic_config["variable"].tolist()
        continuous_vars = continuous_config["variable"].tolist()

        work_df = data_df.copy()
        work_df = coerce_numeric_columns(work_df, item_vars + continuous_vars)

        range_issue_df = build_range_issue_table(work_df, item_vars, args.scale_min, args.scale_max)
        reverse_vars = likert_config.loc[pd.to_numeric(likert_config["reverse"], errors="coerce").fillna(0).astype(int) == 1, "variable"].tolist()
        work_df = reverse_score_items(work_df, reverse_vars, args.scale_min, args.scale_max)

        clean_df, sample_info_df = screen_cases(work_df, item_vars, args.missing_threshold, args.straightline_check)

        item_desc_df = pd.DataFrame()
        if item_vars:
            item_desc_df = descriptive_statistics(clean_df[item_vars], "題項")
            item_desc_df = item_desc_df.merge(ceiling_floor_table(clean_df[item_vars], args.scale_min, args.scale_max), on="題項", how="left")
            item_desc_df["題項標籤"] = item_desc_df["題項"].map(lambda x: label_for(x, label_map))

        construct_score_df = build_construct_scores(clean_df, likert_config)
        construct_desc_df = descriptive_statistics(construct_score_df.drop(columns=["受試者編號"], errors="ignore"), "構面")
        construct_score_df = relabel_construct_outputs(construct_score_df, label_map)
        construct_desc_df = apply_term_labels_to_df(construct_desc_df, ["構面"], label_map)

        alpha_df, item_total_tables, kmo_df, loading_tables, eigen_tables = build_construct_reliability_outputs(clean_df, likert_config)
        alpha_df = apply_term_labels_to_df(alpha_df, ["構面"], label_map)
        kmo_df = apply_term_labels_to_df(kmo_df, ["構面"], label_map)

        relabeled_item_total_tables = {}
        for key, table in item_total_tables.items():
            table = apply_term_labels_to_df(table, ["題項"], label_map)
            relabeled_item_total_tables[label_for(key, label_map)] = table
        relabeled_loading_tables = {}
        relabeled_eigen_tables = {}
        for key, table in loading_tables.items():
            table = apply_term_labels_to_df(table, ["題項"], label_map)
            relabeled_loading_tables[label_for(key, label_map)] = table
        for key, table in eigen_tables.items():
            relabeled_eigen_tables[label_for(key, label_map)] = table

        demographic_df = build_demographic_summary(clean_df, demographic_vars, label_map) if demographic_vars else pd.DataFrame()
        corr_matrix_df, corr_detail_df = build_correlation_results(construct_score_df)
        t_df, anova_df = build_group_difference_results(clean_df, demographic_vars, construct_score_df, label_map)

        warning_df = build_warning_table(item_desc_df, alpha_df, relabeled_item_total_tables, relabeled_loading_tables, range_issue_df)

        analysis_numeric_df = construct_score_df.drop(columns=["受試者編號"], errors="ignore").copy()
        for col in continuous_vars:
            analysis_numeric_df[label_for(col, label_map)] = pd.to_numeric(clean_df[col], errors="coerce")

        if args.models:
            models_df = normalize_models(read_table(Path(args.models)))
        else:
            models_df = pd.DataFrame(columns=["analysis_type", "model_name", "outcome", "predictor", "mediator", "moderator", "covariates", "step_predictors", "notes"])

        # 模型分析使用未標籤化欄位，另建對應資料集
        raw_construct_scores = build_construct_scores(clean_df, likert_config)
        raw_analysis_numeric = raw_construct_scores.drop(columns=["受試者編號"], errors="ignore").copy()
        for col in continuous_vars:
            raw_analysis_numeric[col] = pd.to_numeric(clean_df[col], errors="coerce")

        reg_df, reg_coef_df = run_regression_models(raw_analysis_numeric, models_df, label_map)
        hier_df, hier_coef_df = run_hierarchical_models(raw_analysis_numeric, models_df, label_map)
        med_df, med_detail_df = run_mediation_models(raw_analysis_numeric, models_df, label_map)
        mod_df, mod_coef_df = run_moderation_models(raw_analysis_numeric, models_df, label_map)

        pls_sheets = build_pls_sem_outputs(clean_df, likert_config, raw_construct_scores, label_map)

        apa_demo_df, apa_demo_note = build_apa_demographic_table(demographic_df)
        apa_desc_df, apa_desc_note = build_apa_construct_desc_table(construct_desc_df)
        apa_t_df, apa_t_note = build_apa_ttest_table(t_df)
        apa_a_df, apa_a_note = build_apa_anova_table(anova_df)
        apa_corr_df, apa_corr_note = build_apa_correlation_table(corr_detail_df)
        apa_reg_df, apa_reg_note = build_apa_regression_table(reg_df)
        apa_hier_df, apa_hier_note = build_apa_hierarchical_table(hier_df)
        apa_med_df, apa_med_note = build_apa_mediation_table(med_df)
        apa_mod_df, apa_mod_note = build_apa_moderation_table(mod_df)
        apa_reg_coef_df = build_apa_coeff_table(reg_coef_df)
        apa_hier_coef_df = build_apa_coeff_table(hier_coef_df)
        apa_med_detail_df = build_apa_coeff_table(med_detail_df)
        apa_mod_coef_df = build_apa_coeff_table(mod_coef_df)

        apa_notes_df = pd.DataFrame(
            [
                {"表格": "APA表1_人口統計", "表註": apa_demo_note},
                {"表格": "APA表2_構面描述", "表註": apa_desc_note},
                {"表格": "APA表3_t檢定", "表註": apa_t_note},
                {"表格": "APA表4_ANOVA", "表註": apa_a_note},
                {"表格": "APA表5_相關分析", "表註": apa_corr_note},
                {"表格": "APA表6_多元迴歸", "表註": apa_reg_note},
                {"表格": "APA表7_階層迴歸", "表註": apa_hier_note},
                {"表格": "APA表8_中介分析", "表註": apa_med_note},
                {"表格": "APA表9_調節分析", "表註": apa_mod_note},
            ]
        )

        summary_text = build_summary_text(sample_info_df, demographic_df, alpha_df, kmo_df, corr_detail_df, t_df, anova_df, reg_df, hier_df, med_df, mod_df, warning_df)
        thesis_template_text = build_thesis_paragraphs(summary_text, demographic_df, alpha_df, kmo_df, corr_detail_df, t_df, anova_df, reg_df, hier_df, med_df, mod_df)
        chart_df, chart_message = create_charts(outdir, demographic_df, construct_desc_df, corr_matrix_df)

        project_check_df = build_project_check(Path.cwd())
        output_notes_df = build_output_notes(package_notes, [chart_message])

        excel_sheets: Dict[str, pd.DataFrame] = {
            "專案結構檢查": project_check_df,
            "執行環境備註": output_notes_df,
            "樣本篩選摘要": sample_info_df,
            "人口統計摘要": demographic_df,
            "題項描述統計": item_desc_df,
            "構面分數資料": construct_score_df,
            "構面描述統計": construct_desc_df,
            "信度摘要": alpha_df,
            "KMO_Bartlett": kmo_df,
            "相關分析矩陣": corr_matrix_df,
            "相關分析明細": corr_detail_df,
            "t檢定摘要": t_df,
            "ANOVA摘要": anova_df,
            "多元迴歸摘要": reg_df,
            "多元迴歸係數": reg_coef_df,
            "階層迴歸摘要": hier_df,
            "階層迴歸係數": hier_coef_df,
            "中介分析摘要": med_df,
            "中介分析明細": med_detail_df,
            "調節分析摘要": mod_df,
            "調節分析係數": mod_coef_df,
            "警示訊息": warning_df,
            "APA表1_人口統計": apa_demo_df,
            "APA表2_構面描述": apa_desc_df,
            "APA表3_t檢定": apa_t_df,
            "APA表4_ANOVA": apa_a_df,
            "APA表5_相關分析": apa_corr_df,
            "APA表6_多元迴歸": apa_reg_df,
            "APA表6b_迴歸係數": apa_reg_coef_df,
            "APA表7_階層迴歸": apa_hier_df,
            "APA表7b_階層係數": apa_hier_coef_df,
            "APA表8_中介分析": apa_med_df,
            "APA表8b_中介明細": apa_med_detail_df,
            "APA表9_調節分析": apa_mod_df,
            "APA表9b_調節係數": apa_mod_coef_df,
            "APA表註": apa_notes_df,
            "圖表輸出清單": chart_df,
            "模型設定檔": models_df,
            "變項標籤對照": pd.DataFrame({"variable": list(label_map.keys()), "label": list(label_map.values())}),
        }
        for name, table in relabeled_item_total_tables.items():
            excel_sheets[f"CITC_{name}"] = table
        for name, table in relabeled_loading_tables.items():
            excel_sheets[f"EFA負荷_{name}"] = table
        for name, table in relabeled_eigen_tables.items():
            excel_sheets[f"EFA特徵值_{name}"] = table
        excel_sheets.update(pls_sheets)

        excel_path = outdir / "survey_auto_summary_output.xlsx"
        write_excel(excel_sheets, excel_path)

        thesis_template_path = outdir / "論文段落模板.txt"
        thesis_template_path.write_text(thesis_template_text, encoding=DEFAULT_ENCODING)

        brief_path = outdir / "survey_auto_summary_brief.txt"
        brief_path.write_text(summary_text, encoding=DEFAULT_ENCODING)

        apa_tables_for_word = [
            ("APA 表1 人口統計摘要", apa_demo_df, apa_demo_note),
            ("APA 表2 構面描述統計", apa_desc_df, apa_desc_note),
            ("APA 表3 t 檢定結果", apa_t_df, apa_t_note),
            ("APA 表4 單因子 ANOVA", apa_a_df, apa_a_note),
            ("APA 表5 Pearson 相關分析", apa_corr_df, apa_corr_note),
            ("APA 表6 多元迴歸摘要", apa_reg_df, apa_reg_note),
            ("APA 表7 階層迴歸摘要", apa_hier_df, apa_hier_note),
            ("APA 表8 中介分析摘要", apa_med_df, apa_med_note),
            ("APA 表9 調節分析摘要", apa_mod_df, apa_mod_note),
        ]
        word_path = outdir / "survey_auto_summary_report.docx"
        word_message = write_word_report(word_path, summary_text, apa_tables_for_word, chart_df, package_notes)

        note_df = build_output_notes(package_notes, [chart_message, word_message])
        if not note_df.empty:
            write_excel({"執行環境備註": note_df}, outdir / "survey_auto_summary_notes.xlsx")

        print("分析完成。")
        print(f"Excel 輸出：{excel_path}")
        print(f"Word 輸出：{word_path}")
        print(f"摘要文字：{brief_path}")
        print(f"論文段落模板：{thesis_template_path}")

    except Exception as exc:
        print("執行失敗。")
        print(f"錯誤訊息：{exc}")
        raise


if __name__ == "__main__":
    main()
